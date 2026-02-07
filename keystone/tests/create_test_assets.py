import os
from pathlib import Path
from docx import Document

def create_test_assets():
    data_dir = Path("data/test_data")
    data_dir.mkdir(parents=True, exist_ok=True)
    
    # 1. Create Sample TXT
    txt_path = data_dir / "sample.txt"
    with open(txt_path, "w") as f:
        f.write("Section 1: Introduction\n\nThis is a test paragraph for the document processor.\nIt supports multiple lines.\n\nSection 2: details\n\nAnother paragraph here.")
    print(f"Created {txt_path}")

    # 2. Create Sample DOCX
    docx_path = data_dir / "sample.docx"
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test paragraph in a DOCX file.')
    doc.add_heading('Heading Level 1', level=1)
    doc.add_paragraph('Another paragraph with some bold text.', style='List Bullet')
    doc.core_properties.title = "Test Doc Title"
    doc.core_properties.author = "Test Author"
    doc.save(docx_path)
    print(f"Created {docx_path}")

if __name__ == "__main__":
    create_test_assets()
