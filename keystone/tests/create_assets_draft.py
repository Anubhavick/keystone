import os
from pathlib import Path
from docx import Document
from pypdf import PdfWriter
import io
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

def create_test_assets():
    data_dir = Path("data/test_data")
    data_dir.mkdir(parents=True, exist_ok=True)
    
    # 1. Create Sample TXT
    txt_path = data_dir / "sample.txt"
    with open(txt_path, "w") as f:
        f.write("Section 1: Introduction\n\nThis is a test paragraph for the document processor.\nIt supports multiple lines.\n\nSection 2: details\n\nAnother paragraph here.")
    print(f"Created {txt_path}")

    # 2. Create Sample DOCX
    docx_path = data_dir / "sample.docx"
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test paragraph in a DOCX file.')
    doc.add_heading('Heading Level 1', level=1)
    doc.add_paragraph('Another paragraph with some bold text.', style='List Bullet')
    doc.core_properties.title = "Test Doc Title"
    doc.core_properties.author = "Test Author"
    doc.save(docx_path)
    print(f"Created {docx_path}")

    # 3. Create Sample PDF
    # We use reportlab to generate the PDF content, then pypdf could be used to manipulate if needed, 
    # but reportlab is enough to create. 
    # Wait, reportlab is not in requirements.txt.
    # Let's check requirements.txt content result first.
    # If reportlab is missing, we can use pypdf (which is in requirements) to create a simple PDF?
    # pypdf is mostly for manipulation. Creating from scratch is hard with pypdf alone without an existing PDF.
    # Actually, FPDF is common but not in list. 
    # The requirement list has `pypdf`. 
    # Let's see if we can use a workaround or if `reportlab` is installed in the environment (it's standard in many).
    # If not, I'll use a simple text-based approach for now or mock it in tests.
    # BUT, the prompt asked to "Create small sample files".
    # I will try to use `reportlab` if available. If not, I will mock the PDF extraction test or write a minimal PDF manually.
    # Let's write a minimal PDF structure manually? That's risky.
    # Alternative: Use `pypdf` to merge/write? No, it needs source.
    # Let's assume reportlab might not be there.
    # HOWEVER, checking the prompt "REQUIREMENTS.TXT CONTENT" provided by the user originally:
    # "pypdf==4.0.1", "pdfplumber==0.10.3", "python-docx==1.1.0". 
    # No reportlab.
    # I will stick to testing existing files or mocking for PDF.
    # OR I can try to find a way to create a PDF with just `pypdf`? 
    # Actually, `pypdf` doesn't create text content easily.
    # I'll create the TXT and DOCX. For PDF, I'll create a dummy empty file and rely on mocking in the test, 
    # OR I will try to write a very simple valid PDF header/body manually in python if possible (complex).
    # Best approach: Mock the `pdfplumber.open` in the test so we don't need a real PDF.
    # But the user asked for "Test PDF extraction with sample file".
    # I'll verify if `reportlab` is installed.
    pass

if __name__ == "__main__":
    create_test_assets()
